import io
import json
from datetime import datetime
from app.models.report import Report
from app.models.user import User
from app.models.patient import Patient

# PDF generation
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import (
    getSampleStyleSheet,
    ParagraphStyle
)
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    HRFlowable
)

# DOCX generation
from docx import Document
from docx.shared import (
    Pt,
    RGBColor,
    Cm
)
from docx.enum.text import (
    WD_ALIGN_PARAGRAPH
)


def _flatten(value, indent: int = 0) -> str:
    """
    Recursively convert a nested dict/list into human-readable text.

    dict  → "Key: value" lines (snake_case keys → Title Case)
    list  → bullet lines joined with newlines
    str   → returned as-is (after trying to parse as JSON)
    other → str()
    """
    if value is None:
        return ""

    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return ""
        # Try to parse JSON strings stored by the backend
        if stripped.startswith("{") or stripped.startswith("["):
            try:
                return _flatten(json.loads(stripped), indent)
            except Exception:
                pass
        return stripped

    if isinstance(value, list):
        parts = []
        for item in value:
            flat = _flatten(item, indent)
            if flat:
                parts.append(("  " * indent) + flat)
        return "\n".join(parts)

    if isinstance(value, dict):
        parts = []
        for k, v in value.items():
            if v is None or v == "" or v == [] or v == {}:
                continue
            # snake_case → Title Case label
            label = k.replace("_", " ").title()
            flat_v = _flatten(v, indent + 1)
            if "\n" in flat_v:
                parts.append(("  " * indent) + f"{label}:\n{flat_v}")
            else:
                parts.append(("  " * indent) + f"{label}: {flat_v}")
        return "\n".join(parts)

    return str(value)


def _clean(text) -> str:
    """Return a human-readable string for any SOAP field value."""
    result = _flatten(text)
    return result if result else "N/A"


class ExportService:

    @staticmethod
    def generate_pdf_bytes(
        report: Report,
        doctor: User,
        patient: Patient
    ) -> bytes:

        buffer = io.BytesIO()

        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2 * cm,
            leftMargin=2 * cm,
            topMargin=2 * cm,
            bottomMargin=2 * cm,
        )

        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            "Title",
            parent=styles["Heading1"],
            fontSize=18,
            textColor=colors.HexColor("#1a1a2e"),
            spaceAfter=4,
        )

        heading_style = ParagraphStyle(
            "SectionHead",
            parent=styles["Heading2"],
            fontSize=12,
            textColor=colors.HexColor("#7c3aed"),
            spaceBefore=12,
            spaceAfter=4,
        )

        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontSize=10,
            leading=14,
            textColor=colors.HexColor("#374151"),
        )

        label_style = ParagraphStyle(
            "Label",
            parent=styles["Normal"],
            fontSize=9,
            textColor=colors.HexColor("#6b7280"),
        )

        timestamp = datetime.utcnow().strftime(
            "%B %d, %Y at %H:%M UTC"
        )

        patient_name = _clean(
            getattr(patient, "full_name", None)
            or
            f"{getattr(patient, 'first_name', '')} "
            f"{getattr(patient, 'last_name', '')}".strip()
        )

        doctor_name = _clean(
            getattr(doctor, "full_name", None)
        )

        license_no = _clean(
            getattr(doctor, "license_number", None)
        )

        story = []

        # Header
        story.append(
            Paragraph(
                "MediScribe",
                title_style
            )
        )

        story.append(
            Paragraph(
                "Clinical SOAP Report",
                styles["Heading2"]
            )
        )

        story.append(
            HRFlowable(
                width="100%",
                thickness=1,
                color=colors.HexColor("#e5e7eb")
            )
        )

        story.append(
            Spacer(1, 0.3 * cm)
        )

        # Meta table
        meta_data = [
            [
                "Report ID",
                report.report_id[:8].upper() + "..."
            ],
            [
                "Status",
                (report.status or "draft").capitalize()
            ],
            [
                "Generated",
                timestamp
            ],
            [
                "Doctor",
                f"Dr. {doctor_name}  |  License: {license_no}"
            ],
            [
                "Patient",
                patient_name
            ],
        ]

        meta_table = Table(
            meta_data,
            colWidths=[4 * cm, 13 * cm]
        )

        meta_table.setStyle(
            TableStyle([
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TEXTCOLOR", (0, 0), (0, -1),
                 colors.HexColor("#6b7280")),
                ("TEXTCOLOR", (1, 0), (1, -1),
                 colors.HexColor("#111827")),
                ("FONTNAME", (0, 0), (0, -1),
                 "Helvetica-Bold"),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
            ])
        )

        story.append(meta_table)

        story.append(
            Spacer(1, 0.4 * cm)
        )

        story.append(
            HRFlowable(
                width="100%",
                thickness=0.5,
                color=colors.HexColor("#e5e7eb")
            )
        )

        # SOAP sections
        sections = [
            ("S — Subjective", report.subjective),
            ("O — Objective", report.objective),
            ("A — Assessment", report.assessment),
            ("P — Plan", report.plan),
        ]

        for title, content in sections:

            story.append(
                Paragraph(
                    title,
                    heading_style
                )
            )

            text = _clean(content)

            text = (
                text
                .replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
                .replace("\n", "<br/>")
            )

            story.append(
                Paragraph(
                    text,
                    body_style
                )
            )

            story.append(
                Spacer(1, 0.2 * cm)
            )

        # Medications
        meds = report.medications

        if meds:

            story.append(
                Paragraph(
                    "Medications",
                    heading_style
                )
            )

            if isinstance(meds, list):

                for med in meds:

                    if isinstance(med, dict):

                        name = med.get(
                            "name",
                            med.get(
                                "drug",
                                str(med)
                            )
                        )

                        dose = med.get(
                            "dose",
                            med.get(
                                "dosage",
                                ""
                            )
                        )

                        freq = med.get(
                            "frequency",
                            med.get(
                                "freq",
                                ""
                            )
                        )

                        line = f"• {name}"

                        if dose:
                            line += f"  {dose}"

                        if freq:
                            line += f"  ({freq})"

                        story.append(
                            Paragraph(
                                line,
                                body_style
                            )
                        )

                    else:

                        story.append(
                            Paragraph(
                                f"• {med}",
                                body_style
                            )
                        )

            else:

                story.append(
                    Paragraph(
                        _clean(meds),
                        body_style
                    )
                )

        # Follow-up
        if report.follow_up_needed:

            story.append(
                Paragraph(
                    "Follow-Up",
                    heading_style
                )
            )

            days = report.follow_up_days

            fu_text = (
                f"Follow-up required in {days} day(s)."
                if days
                else "Follow-up required."
            )

            story.append(
                Paragraph(
                    fu_text,
                    body_style
                )
            )

        # Footer
        story.append(
            Spacer(1, 0.5 * cm)
        )

        story.append(
            HRFlowable(
                width="100%",
                thickness=0.5,
                color=colors.HexColor("#e5e7eb")
            )
        )

        story.append(
            Spacer(1, 0.2 * cm)
        )

        story.append(
            Paragraph(
                f"Digitally prepared by MediScribe AI Platform • {timestamp}",
                label_style
            )
        )

        doc.build(story)

        buffer.seek(0)

        return buffer.read()

    @staticmethod
    def generate_docx_bytes(
        report: Report,
        doctor: User,
        patient: Patient
    ) -> bytes:

        doc = Document()

        for section in doc.sections:

            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        timestamp = datetime.utcnow().strftime(
            "%B %d, %Y at %H:%M UTC"
        )

        patient_name = _clean(
            getattr(patient, "full_name", None)
            or
            f"{getattr(patient, 'first_name', '')} "
            f"{getattr(patient, 'last_name', '')}".strip()
        )

        doctor_name = _clean(
            getattr(doctor, "full_name", None)
        )

        license_no = _clean(
            getattr(doctor, "license_number", None)
        )

        title = doc.add_heading(
            "MediScribe — Clinical SOAP Report",
            0
        )

        title.alignment = (
            WD_ALIGN_PARAGRAPH.LEFT
        )

        title.runs[0].font.color.rgb = RGBColor(
            0x1a,
            0x1a,
            0x2e
        )

        doc.add_paragraph()

        # SOAP sections
        sections = [
            ("S — Subjective", report.subjective),
            ("O — Objective", report.objective),
            ("A — Assessment", report.assessment),
            ("P — Plan", report.plan),
        ]

        for title_text, content in sections:

            h = doc.add_heading(
                title_text,
                level=2
            )

            h.runs[0].font.color.rgb = RGBColor(
                0x7c,
                0x3a,
                0xed
            )

            text = _clean(content)

            doc.add_paragraph(text)

        buffer = io.BytesIO()

        doc.save(buffer)

        buffer.seek(0)

        return buffer.read()